from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Pt


# ================= PDF EXPORT =================
def generate_lesson_pdf(lesson, school_name=""):
    file_path = "lesson_plan.pdf"

    doc = SimpleDocTemplate(file_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Header
    title = f"<b>{school_name}</b><br/>Class {lesson['grade']} – {lesson['subject']}<br/>"
    title += f"{lesson['chapter']} (Day {lesson['day_no']} of {lesson['total_days']})"
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))

    # Pedagogy
    story.append(Paragraph(
        f"<b>Pedagogy:</b> {lesson['pedagogy']} | "
        f"<b>Integration:</b> {lesson['integration'] or '—'}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 12))

    # Learning Outcomes
    story.append(Paragraph("<b>Learning Outcomes</b>", styles["Heading2"]))
    for lo in lesson["learning_outcomes"]:
        story.append(Paragraph(f"- {lo}", styles["Normal"]))
    story.append(Spacer(1, 12))

    # Lesson Table
    table_data = [["Phase", "Time", "Focus"]]
    for phase, mins in lesson["period_structure"]:
        table_data.append([phase, f"{mins} min", "Teaching & Interaction"])

    table = Table(table_data, colWidths=[120, 80, 280])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#E3F2FD")),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#000000")),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
    ]))

    story.append(table)
    story.append(Spacer(1, 12))

    # Teaching Script
    story.append(Paragraph("<b>Detailed Teaching Script</b>", styles["Heading2"]))
    story.append(Paragraph(lesson["script"].replace("\n", "<br/>"), styles["Normal"]))

    doc.build(story)
    return file_path


# ================= WORD EXPORT =================
def generate_lesson_word(lesson, school_name=""):
    file_path = "lesson_plan.docx"
    doc = Document()

    # Title
    title = doc.add_heading(level=1)
    run = title.add_run(school_name)
    run.bold = True

    doc.add_paragraph(
        f"Class {lesson['grade']} – {lesson['subject']}\n"
        f"{lesson['chapter']} (Day {lesson['day_no']} of {lesson['total_days']})"
    )

    # Pedagogy
    p = doc.add_paragraph()
    p.add_run("Pedagogy: ").bold = True
    p.add_run(lesson["pedagogy"])
    p.add_run(" | Integration: ").bold = True
    p.add_run(lesson["integration"] or "—")

    # Learning Outcomes
    doc.add_heading("Learning Outcomes", level=2)
    for lo in lesson["learning_outcomes"]:
        doc.add_paragraph(lo, style="List Bullet")

    # Lesson Table
    doc.add_heading("Lesson Flow", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Phase"
    hdr[1].text = "Time"
    hdr[2].text = "Focus"

    for phase, mins in lesson["period_structure"]:
        row = table.add_row().cells
        row[0].text = phase
        row[1].text = f"{mins} min"
        row[2].text = "Teaching & Interaction"

    # Script
    doc.add_heading("Detailed Teaching Script", level=2)
    doc.add_paragraph(lesson["script"])

    doc.save(file_path)
    return file_path
