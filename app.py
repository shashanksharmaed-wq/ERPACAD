import streamlit as st
import pandas as pd
from openai import OpenAI
import math
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# -----------------------------------------------------------------------------
# 1. APP CONFIGURATION
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="Teach+ Pro", 
    page_icon="🎓", 
    layout="wide"
)

st.title("🎓 Teach+ Curriculum Director")
st.markdown("### Strategic Pacing & Micro-Planning System")

# -----------------------------------------------------------------------------
# 2. LOAD DATABASE
# -----------------------------------------------------------------------------
@st.cache_data
def load_data():
    file_path = "Teachshank_Master_Database_FINAL.tsv"
    try:
        df = pd.read_csv(file_path, sep='\t')
        return df
    except FileNotFoundError:
        st.error(f"⚠️ Critical Error: Could not find '{file_path}'. Please upload it to your GitHub repository.")
        return pd.DataFrame()

df = load_data()
if df.empty:
    st.stop()

# -----------------------------------------------------------------------------
# 3. SIDEBAR: GLOBAL SELECTORS
# -----------------------------------------------------------------------------
st.sidebar.header("1️⃣ Global Context")

# Grade & Subject Selection
grade_list = sorted(df['Grade'].astype(str).unique().tolist())
selected_grade = st.sidebar.selectbox("Select Grade", grade_list)

subject_list = sorted(df[df['Grade'].astype(str) == selected_grade]['Subject'].unique().tolist())
selected_subject = st.sidebar.selectbox("Select Subject", subject_list)

# Filter Data for this specific Class & Subject
subject_data = df[
    (df['Grade'].astype(str) == selected_grade) & 
    (df['Subject'] == selected_subject)
]

# -----------------------------------------------------------------------------
# 4. TAB SYSTEM (The Fix for Visibility)
# -----------------------------------------------------------------------------
# We create two distinct tabs to separate the "Year View" from the "Day View"
tab1, tab2 = st.tabs(["📅 Annual Pacing Calendar", "📝 Daily Lesson Designer"])

# =============================================================================
# TAB 1: ANNUAL CALENDAR VIEW
# =============================================================================
with tab1:
    st.subheader(f"Annual Roadmap: {selected_grade} - {selected_subject}")
    st.markdown("This calendar calculates the **Ideal Pacing** (180 Days) based on the number of learning outcomes per chapter.")
    
    if not subject_data.empty:
        # --- CALCULATION LOGIC ---
        # 1. Count outcomes per chapter
        chapter_groups = subject_data.groupby('Chapter Name')['Learning Outcomes'].count().reset_index()
        chapter_groups.columns = ['Unit / Chapter Name', 'Topic_Count']
        
        # 2. Calculate Total Outcomes for the subject
        total_outcomes = chapter_groups['Topic_Count'].sum()
        
        if total_outcomes > 0:
            # 3. Distribute 180 days proportionally
            # Formula: (Topics in Chapter / Total Topics) * 180
            chapter_groups['Allocated Days'] = (chapter_groups['Topic_Count'] / total_outcomes) * 180
            chapter_groups['Allocated Days'] = chapter_groups['Allocated Days'].apply(lambda x: math.ceil(x))
            
            # 4. Display as a clean Table
            st.dataframe(
                chapter_groups,
                use_container_width=True,
                height=500,
                column_config={
                    "Unit / Chapter Name": st.column_config.TextColumn("Instructional Unit", width="large"),
                    "Topic_Count": st.column_config.NumberColumn("Topics", width="small"),
                    "Allocated Days": st.column_config.ProgressColumn(
                        "Pacing (Days)", 
                        format="%d Days", 
                        min_value=0, 
                        max_value=chapter_groups['Allocated Days'].max()
                    )
                }
            )
            
            # Summary Metrics
            c1, c2, c3 = st.columns(3)
            c1.metric("Total Instructional Days", "180")
            c2.metric("Total Topics/Outcomes", total_outcomes)
            c3.metric("Avg Days per Topic", f"{180/total_outcomes:.1f}")
            
        else:
            st.warning("No outcomes found to generate calendar.")
    else:
        st.info("Please select a Grade and Subject to view the calendar.")


# =============================================================================
# TAB 2: DAILY LESSON DESIGNER
# =============================================================================
with tab2:
    st.subheader("Micro-Lesson Generator")
    
    col_a, col_b = st.columns([1, 2])
    
    with col_a:
        st.markdown("**Step 1: Lesson Context**")
        # Chapter Selector
        chapter_list = subject_data['Chapter Name'].unique().tolist()
        selected_chapter = st.selectbox("Select Lesson / Chapter", chapter_list)

        # Outcome Selector
        outcome_list = subject_data[subject_data['Chapter Name'] == selected_chapter]['Learning Outcomes'].dropna().tolist()
        if outcome_list:
            selected_outcome = st.selectbox("Select Specific Learning Focus", outcome_list)
        else:
            selected_outcome = "General Chapter Overview"

        st.markdown("---")
        st.markdown("**Step 2: Configuration**")
        
        # Duration Slider
        duration = st.slider("Duration (Mins)", 30, 90, 45)
        
        # Phase Selector
        lesson_phase = st.radio(
            "Lesson Phase:",
            ["Introduction / Discovery", "Deep Dive / Explanation", "Practice / Application", "Assessment"],
            index=0
        )
        
        generate_btn = st.button("🚀 Generate Lesson Plan", type="primary")

    with col_b:
        # -----------------------------------------------------------------------------
        # DOCX GENERATOR FUNCTION
        # -----------------------------------------------------------------------------
        def create_word_docx(lesson_text, metadata):
            doc = Document()
            title = doc.add_heading(f"Teach+ Lesson Plan", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata Table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            p1 = hdr_cells[0].paragraphs[0]
            p1.add_run(f"Class: {metadata['grade']} | Subject: {metadata['subject']}\n").bold = True
            p1.add_run(f"Unit: {metadata['chapter']}")
            
            p2 = hdr_cells[1].paragraphs[0]
            p2.add_run(f"Duration: {metadata['duration']} Mins | Phase: {metadata['phase']}\n").bold = True
            p2.add_run(f"Topic: {metadata['topic']}")
            
            doc.add_paragraph("") 

            # Content Parser
            lines = lesson_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line: continue
                if line.startswith('## '):
                    h = doc.add_heading(line.replace('## ', ''), level=1)
                    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                elif line.startswith('### '):
                    h = doc.add_heading(line.replace('### ', ''), level=2)
                    h.runs[0].font.color.rgb = RGBColor(51, 102, 153)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line.replace('- ', '').replace('* ', ''), style='List Bullet')
                else:
                    doc.add_paragraph(line)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        # -----------------------------------------------------------------------------
        # AI ENGINE
        # -----------------------------------------------------------------------------
        def generate_lesson_script(grade, subject, chapter, outcome, duration, phase):
            # Timeboxing Logic
            t_engage = int(duration * 0.15)
            t_explore = int(duration * 0.30)
            t_explain = int(duration * 0.25)
            
            prompt = f"""
            You are the 'Teach+' Lesson Architect.
            
            TASK:
            Generate a deep, scripted lesson plan for:
            - **Grade:** {grade}
            - **Subject:** {subject}
            - **Lesson Unit:** {chapter}
            - **Focus Topic:** {outcome}
            - **Phase:** {phase}
            - **Duration:** {duration} Minutes
            
            NOTE:
            The duration is strictly {duration} minutes. Adjust the depth of activities to fit this time.
            If 30 mins -> Keep it sharp and focused.
            If 90 mins -> Include deeper exploration and more student practice time.

            OUTPUT FORMAT (Markdown):

            ## 1. Context
            - **Topic:** (Specific Focus)
            - **Goal:** (What specifically will be mastered?)

            ## 2. The Micro-Script ({duration} Mins)

            ### ⏰ 00:00 - 00:{t_engage:02d} | I. ENGAGE
            - **Script:** "Teacher says..."
            - **Action:** (Prop/Movement)

            ### ⏰ 00:{t_engage:02d} - 00:{t_engage+t_explore:02d} | II. EXPLORE
            - **Activity:** ...
            - **Scripted Instruction:** ...

            ### ⏰ 00:{t_engage+t_explore:02d} - 00:{t_engage+t_explore+t_explain:02d} | III. EXPLAIN
            - **Concept:** ...
            - **Script:** ...

            ### ⏰ 00:{t_engage+t_explore+t_explain:02d} - 00:{duration:02d} | IV. ELABORATE & EVALUATE
            - **Application Task:** ...
            - **Exit Ticket:** ...
            """
            
            try:
                client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
                response = client.chat.completions.create(
                    model="gpt-4o", # Use gpt-3.5-turbo if budget is tight
                    messages=[
                        {"role": "system", "content": "You are an expert pedagogue."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7
                )
                return response.choices[0].message.content
            except Exception as e:
                return f"Error: {str(e)}"

        # -----------------------------------------------------------------------------
        # EXECUTION
        # -----------------------------------------------------------------------------
        if generate_btn:
            if "OPENAI_API_KEY" not in st.secrets:
                st.error("⚠️ OpenAI API Key is missing in Streamlit Secrets!")
            else:
                with st.spinner(f"Designing Lesson..."):
                    plan_text = generate_lesson_script(
                        selected_grade, 
                        selected_subject, 
                        selected_chapter, 
                        selected_outcome, 
                        duration,
                        lesson_phase
                    )
                    
                    st.success("Plan Generated Successfully!")
                    st.markdown(plan_text)
                    
                    # Word Doc Download
                    metadata = {
                        "grade": selected_grade,
                        "subject": selected_subject,
                        "chapter": selected_chapter,
                        "topic": selected_outcome,
                        "duration": duration,
                        "phase": lesson_phase
                    }
                    docx_file = create_word_docx(plan_text, metadata)
                    safe_chapter = "".join([c for c in selected_chapter if c.isalnum() or c in (' ','-')]).strip()
                    
                    st.download_button(
                        label="📄 Download as Word Doc (.docx)",
                        data=docx_file,
                        file_name=f"TeachPlus_{safe_chapter}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
